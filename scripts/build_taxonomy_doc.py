import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATA_PATH = ROOT / "docs" / "taxonomy-data.json"
OUTPUT_PATH = ROOT / "docs" / "Dimensions_of_Expression_Taxonomy.docx"

BLUE = "246A73"
DARK_BLUE = "1F4D78"
MUTED = "657078"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
LINE = "D9DDD9"
WHITE = "FFFFFF"
INK = "202124"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total = sum(widths_dxa)
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_font(run, name="Calibri", size=11, color=INK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def add_term_table(doc, rows, first_label="Category", widths=(1800, 7020, 540)):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = [first_label, "Tags", "Count"]
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        set_font(run, size=9.5, color=DARK_BLUE, bold=True)
        if index == 2:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for label, terms in rows:
        cells = table.add_row().cells
        label_paragraph = cells[0].paragraphs[0]
        label_paragraph.paragraph_format.space_after = Pt(0)
        run = label_paragraph.add_run(label)
        set_font(run, size=9.5, color=INK, bold=True)

        terms_paragraph = cells[1].paragraphs[0]
        terms_paragraph.paragraph_format.space_after = Pt(0)
        run = terms_paragraph.add_run(", ".join(terms))
        set_font(run, size=9.5, color=INK)

        count_paragraph = cells[2].paragraphs[0]
        count_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        count_paragraph.paragraph_format.space_after = Pt(0)
        run = count_paragraph.add_run(str(len(terms)))
        set_font(run, size=9.5, color=MUTED, bold=True)

    set_table_geometry(table, list(widths))
    set_repeat_table_header(table.rows[0])
    for row in table.rows[1:]:
        prevent_row_split(row)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_field_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for index, header in enumerate(("Element", "Purpose")):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        run = cell.paragraphs[0].add_run(header)
        set_font(run, size=9.5, color=DARK_BLUE, bold=True)
    for item in rows:
        cells = table.add_row().cells
        run = cells[0].paragraphs[0].add_run(item["name"])
        set_font(run, size=9.5, bold=True)
        run = cells[1].paragraphs[0].add_run(item["purpose"])
        set_font(run, size=9.5)
    set_table_geometry(table, [2700, 6660])
    set_repeat_table_header(table.rows[0])
    for row in table.rows[1:]:
        prevent_row_split(row)
    return table


def set_keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_section_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    set_keep_with_next(paragraph)
    return paragraph


def total_terms(groups):
    return sum(len(group["terms"]) for group in groups)


def build_document(data):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("DIMENSIONS OF EXPRESSION  |  TAXONOMY REFERENCE")
    set_font(run, size=8.5, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(2)
    run = kicker.add_run("CURRENT APPLICATION VOCABULARY")
    set_font(run, size=9, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(5)
    run = title.add_run("Dimensions of Expression Taxonomy")
    set_font(run, size=24, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run(
        "A structured inventory of every vocabulary tag, emotion option, feature descriptor, and coding field currently available in the application."
    )
    set_font(run, size=11.5, color=MUTED)

    vocab_count = sum(total_terms(category["groups"]) for category in data["taxonomy"])
    feature_count = sum(len(feature["terms"]) for feature in data["featuresActivated"])
    emotion_count = sum(len(family["terms"]) for family in data["emotionFamilies"])
    event_count = sum(len(dimension["terms"]) for dimension in data["whatJustHappenedDimensions"])
    unique_terms = set()
    for category in data["taxonomy"]:
        for group in category["groups"]:
            unique_terms.update(term.lower() for term in group["terms"])
    for feature in data["featuresActivated"]:
        unique_terms.update(term.lower() for term in feature["terms"])
    for family in data["emotionFamilies"]:
        unique_terms.update(term.lower() for term in family["terms"])
    for dimension in data["whatJustHappenedDimensions"]:
        unique_terms.update(term.lower() for term in dimension["terms"])

    summary = doc.add_table(rows=2, cols=5)
    summary.style = "Table Grid"
    summary_values = [
        ("Vocabulary tags", str(vocab_count)),
        ("Feature descriptors", str(feature_count)),
        ("Emotion options", str(emotion_count)),
        ("Event options", str(event_count)),
        ("Unique terms", str(len(unique_terms))),
    ]
    for index, (label, value) in enumerate(summary_values):
        top = summary.rows[0].cells[index]
        bottom = summary.rows[1].cells[index]
        set_cell_shading(top, LIGHT_BLUE)
        set_cell_shading(bottom, WHITE)
        top_p = top.paragraphs[0]
        top_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        top_p.paragraph_format.space_after = Pt(0)
        run = top_p.add_run(label)
        set_font(run, size=8.5, color=DARK_BLUE, bold=True)
        bottom_p = bottom.paragraphs[0]
        bottom_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bottom_p.paragraph_format.space_after = Pt(0)
        run = bottom_p.add_run(value)
        set_font(run, size=16, color=BLUE, bold=True)
    set_table_geometry(summary, [1872, 1872, 1872, 1872, 1872])
    set_repeat_table_header(summary.rows[0])
    prevent_row_split(summary.rows[1])

    add_section_heading(doc, "How to read the hierarchy", 1)
    hierarchy = [
        ("Workflow area", "The coding step or conceptual layer, such as Vocabulary, Emotion, or Features Activated."),
        ("Umbrella category", "The broad organizing family, such as Appraisal, Joy, or Brows."),
        ("Subgroup", "A narrower conceptual branch, such as Time direction, Agitation, or Fearful."),
        ("Tag", "The selectable term saved in an External or Internal bin, emotion slider, or feature descriptor list."),
    ]
    add_term_table(doc, [(label, [detail]) for label, detail in hierarchy], "Level")

    add_section_heading(doc, "1. Features Activated", 1)
    paragraph = doc.add_paragraph(
        "Each anatomical or postural feature has an intensity slider from 0 to 100 and a multi-select descriptor list."
    )
    paragraph.paragraph_format.space_after = Pt(8)
    add_term_table(
        doc,
        [(feature["name"], feature["terms"]) for feature in data["featuresActivated"]],
        "Feature",
    )

    add_section_heading(doc, "2. Emotion Blend Sliders", 1)
    paragraph = doc.add_paragraph(
        "The application provides two configurable sliders for each umbrella emotion. The dropdown label can be changed to any term within that family."
    )
    paragraph.paragraph_format.space_after = Pt(8)
    emotion_rows = []
    for family in data["emotionFamilies"]:
        defaults = " / ".join(family["defaults"])
        emotion_rows.append((f'{family["name"]} (defaults: {defaults})', family["terms"]))
    add_term_table(doc, emotion_rows, "Umbrella")

    aliases = [
        (alias.title(), [family])
        for alias, family in sorted(data["emotionFamilyAliases"].items())
    ]
    add_section_heading(doc, "Consensus mapping aliases", 2)
    doc.add_paragraph(
        "These legacy or broad saved labels are mapped to the five umbrella families when the consensus radar is calculated."
    )
    add_term_table(doc, aliases, "Saved label", widths=(2200, 6620, 540))

    add_section_heading(doc, "3. What Just Happened", 1)
    paragraph = doc.add_paragraph(
        "Each dimension uses a selectable interpretation label and a 0 to 100 intensity slider to code the inferred change process."
    )
    paragraph.paragraph_format.space_after = Pt(8)
    event_rows = []
    for dimension in data["whatJustHappenedDimensions"]:
        event_rows.append(
            (
                f'{dimension["name"]} (default: {dimension["defaultTerm"]})',
                dimension["terms"],
            )
        )
    add_term_table(doc, event_rows, "Dimension")

    add_section_heading(doc, "4. Vocabulary Assignment Taxonomy", 1)
    doc.add_paragraph(
        "Every tag in this section can be assigned to either the External or Internal bin. Repeated words are retained where they carry different meanings in different branches."
    )
    for category in data["taxonomy"]:
        add_section_heading(doc, category["title"], 2)
        meta = doc.add_paragraph()
        meta.paragraph_format.space_after = Pt(7)
        source_run = meta.add_run(f'Source frame: {category["source"]}. ')
        set_font(source_run, size=9.5, color=MUTED, bold=True)
        desc_run = meta.add_run(category["description"])
        set_font(desc_run, size=9.5, color=MUTED)
        add_term_table(
            doc,
            [(group["label"], group["terms"]) for group in category["groups"]],
            "Subgroup",
        )

    add_section_heading(doc, "5. Assignment Bins and Interpretation Fields", 1)
    add_section_heading(doc, "Message bins", 2)
    add_field_table(doc, data["assignmentBins"])
    doc.add_paragraph()
    add_section_heading(doc, "Interpretation fields", 2)
    add_field_table(doc, data["interpretationFields"])

    add_section_heading(doc, "6. Workflow Modules", 1)
    module_rows = [(module["label"], [module["id"]]) for module in data["workflowModules"]]
    add_term_table(doc, module_rows, "Displayed label", widths=(2600, 6220, 540))

    add_section_heading(doc, "7. Literature Backbone Notes", 1)
    for note in data["literatureNotes"]:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        run = paragraph.add_run(note)
        set_font(run, size=10.5)

    return doc


def main():
    if len(sys.argv) > 1:
        data_path = Path(sys.argv[1])
    else:
        data_path = DATA_PATH
    if len(sys.argv) > 2:
        output_path = Path(sys.argv[2])
    else:
        output_path = OUTPUT_PATH

    data = json.loads(data_path.read_text())
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = build_document(data)
    document.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
